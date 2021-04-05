import json
import csv
from docx.oxml.shared import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.oxml import CT_P
from datetime import datetime
import locale
import xlrd
from openpyxl import load_workbook


def get_data_from_json(folder): # load the config.json file
    with open('{}/config.json'.format(folder)) as filePath:
        return json.load(filePath)

# defined function to get all site name
def atlas_list(data):
    allAtlas = []
    for atlas in data['Referenzen']:
        if atlas['type'] == 'Atlas':
            allAtlas.append(atlas['name'])
        if atlas['name'] == 'ConWx':
            allAtlas.append('ConWx')
    return allAtlas

# format the name of the sites
#defined function to format the name of the atlas
def atlasName(data):
    newS = {}
    for s in atlas_list(data):
        tempS = []
        for char in s:
            if char == '-':
                continue
            if char == '.':
                tempS.append('_')
                continue
            tempS.append(char)
        newS[s] = ''.join(tempS)
    return newS

#define a function to get the cell string of atlas
def cell_string(data):
    cell_string_list = {}
    for atlas in atlas_list(data):
        for cell in data['selected cells']:
            if atlas == cell:
                cell_string_list[cell] = data['selected cells'][cell]
                #print(cell)
    return cell_string_list

#height associated with index number and atlas name, passed two parameters
def height_list(data, index, atlas):
    allAtlas = []
    height_string_list = 0
    for atlasI in data["Referenzen"]:
        allAtlas.append(atlasI['name'])
    for atlas_length in range(len(allAtlas)):      #loop through all atlases
        if allAtlas[atlas_length] == atlas:
            height_string_list = data['Referenzen'][atlas_length]['heights'][index]
            break
    return height_string_list

#defined a function of refprd for all the atlas
def refprd_list(data,atlas):
    allAtlas = []
    refprd_string_list = ''
    for atlasI in data['Referenzen']:
        allAtlas.append(atlasI['name'])
    for atlas_length in range(len(allAtlas)):
        if allAtlas[atlas_length] == atlas:
            refprd_string_list = data['Referenzen'][atlas_length]['refprd']
            break
    return refprd_string_list

def get_bookmark_list(doc):
    """
    Return list of bookmarks in a document
    """
    bookmarkList = []
    doc_element = doc.part.element.body
    #bookmarks_list = doc_element.findall('.//' + qn('w:bookmarkStart'))
    #bookmarks_list = doc_element.findall('.//' + qn('w:bookmarkStart'))

    bookmarks_list = doc_element.findall('w:bookmarkStart')
    print(bookmarks_list)
    for bookmark in bookmarks_list:
        bookmarkList.append(bookmark.get(qn('w:name')))
    return bookmarkList

def print_bookmark_list(doc):
    bookmark_list = get_bookmark_list(doc)
    print("Here are the list of the bookmark in the document:")
    for i in range(len(bookmark_list)):
        print(bookmark_list[i])

def make_rows_italic(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in paragraph.runs:
                    run.font.italic = True


def bookmark_text(doc, bookmark_name, text, underline = False, italic = False, bold = False, style = None):
    doc_element = doc._part._element
    bookmarks_list = doc_element.findall('.//' + qn('w:bookmarkStart'))
    for bookmark in bookmarks_list:
        name = bookmark.get(qn('w:name'))
        if name == bookmark_name:
            par = bookmark.getparent()
            if not isinstance(par, CT_P):
                return False
            else:
                i = par.index(bookmark) + 1
                p = doc.add_paragraph()
                run = p.add_run(text, style)
                run.underline = underline
                run.italic = italic
                run.bold = bold
                par.insert(i,run._element)
                p = p._element
                p.getparent().remove(p)
                p._p = p._element = None
                return True
    return False

def count_missing_values(folderIn,refName,data,wtgnum):
    '''This will loop through the csv file to check if the value is -9.0 and increment the counter'''
    countM = 0
    with open("{}/{}/{}".format(folderIn, refName, data['WEAs'][wtgnum - 1]['fname']), newline='') as csvfile:
        reader = csv.DictReader(csvfile)
        for row in reader:
            if row["WEA {}".format(wtgnum)] == '-9.0':
                countM += 1
    return countM

def bookmark_picture(doc, bookmark_name, path):
    doc_element = doc._part._element
    bookmarks_list = doc_element.findall('.//' + qn('w:bookmarkStart'))
    for bookmark in bookmarks_list:
        name = bookmark.get(qn('w:name'))
        if name == bookmark_name:
            par = bookmark.getparent()
            if not isinstance(par, CT_P):
                return False
            else:
                i = par.index(bookmark) + 1
                p = doc.add_paragraph()
                run = p.add_run()
                run.add_picture(path, width=Inches(2.60), height=Inches(1.70))
                par.insert(i, run._element)
                p = p._element
                p.getparent().remove(p)
                p._p = p._element = None
                return True
    return False

def get_date(folderIn, refName, data, English, weaNo):
    '''to read csv file we need csv library. it has DictReader module which helps reading the file.
        we read each entry of a column. then we convert the first two dates from string to datetime object. then
        we check the difference between to date to find out in which format we will insert the dates in the doc.'''
    dateList = []
    with open("{}/{}/{}".format(folderIn, refName, data['WEAs'][weaNo - 1]['fname']), newline='') as csvfile:
        reader = csv.DictReader(csvfile)
        for row in reader:
            dateList.append(row['Datum'])

    senTimeStr = ''
    d_1st_str = ''
    d_last_str = ''

    if data['WEAs'][weaNo - 1]['timestep'] == 'monthly':
        if English == True:
            senTimeStr = 'monthly production values are missing or are interpreted as missing values. Availability values are present.'
            d_1st_str = dateList[0][0:3] + dateList[0][6:10]
            d_last_str = dateList[-1][0:3] + dateList[-1][6:10]

        if English == False:
            senTimeStr = 'Monatserträge fehlen oder werden als Fehlwerte angesehen. Verfügbarkeitsangaben sind vorhanden.'
            d_1st = datetime.strptime(dateList[0], '%m/%d/%Y').date()
            d_last = datetime.strptime(dateList[-1], '%m/%d/%Y').date()
            # locale is used to make regional changes
            locale.setlocale(locale.LC_TIME, 'de_DE.UTF-8')
            d_1st_str = d_1st.strftime('%b. %Y')
            d_last_str = d_last.strftime('%b. %Y')
    elif data['WEAs'][weaNo - 1]['timestep'] == 'daily':
        if English == True:
            senTimeStr = 'daily production values are missing or are interpreted as missing values. Availability values are present.'
            d_1st_str = dateList[0]
            d_last_str = dateList[-1]
        if English == False:
            senTimeStr = 'Tagesertrag fehlen oder werden als Fehlwerte angesehen. Verfügbarkeitsangaben sind vorhanden.'
            d_1st = datetime.strptime(dateList[0], '%m/%d/%Y').date()
            d_last = datetime.strptime(dateList[-1], '%m/%d/%Y').date()
            locale.setlocale(locale.LC_TIME, 'de_DE.UTF-8')
            d_1st_str = d_1st.strftime('%-d %b. %Y')  # removed the leading zero. 1.12.1999
            d_last_str = d_last.strftime('%-d %b. %Y')

    d_all_str = []
    d_all_str.append(d_1st_str)
    d_all_str.append(d_last_str)
    d_all_str.append(senTimeStr)

    return d_all_str

def get_coordinates(folderIn, refName, weaNo, data, English):
    '''First it will read the csv file. Then it will store all its content in a list (corlist). Then
        another list will be created to get the value without the tab (tempcorlist). Then we use a dictionary
        where value of the 2nd column of the csv file will be the key of dictionary and value of the 1st
        column will be the value of dictionary.'''
    corList = []
    with open("{}/{}/wea{}-cor.csv".format(folderIn, refName, weaNo), newline='') as csvfile:
        reader = csv.reader(csvfile)
        for row in reader:
            corList.append(row)
    tempCorList = []
    newCorList = {}

    i = 0
    for cor in corList:
        tempCorList.append(corList[i][0].split('\t'))
        i += 1

    i = 0
    for cor in corList:
        newCorList[tempCorList[i][1]] = tempCorList[i][0]
        i += 1

    # create appropriate key to get proper coordinates value from cfg file
    firstSplit = newCorList[cell_string(data)[refName]].split('\t')
    secondSplit = firstSplit[0].split('x')
    val1 = secondSplit[1]
    val2 = secondSplit[0]
    if val2[0] == '-':
        newVal2 = val2[1:-1]
        cor1 = "{:.2f}".format(float(val1))
        cor2 = "{:.2f}".format(float(newVal2))
        addVar1 = "{}°S {}°W".format(cor1, cor2)
    else:
        cor1 = "{:.2f}".format(float(val1))
        cor2 = "{:.2f}".format(float(val2))
        addVar1 = "{}°N {}°E".format(cor1, cor2)

    addVar2 = height_list(data, int(cell_string(data)[refName][-1]), refName)
    addVar3 = refName

    sentence = ''
    if English == True:
        sentence = "The appropriate cell is chosen via regression analysis. The best correlated cell is used for the long-term correlation. " \
                    "The coordinates of the chosen cell are: {}, {} m over ground. The {}-Index data are adapted to the measurement time series via " \
                    "a regression equation.".format(addVar1, addVar2, addVar3)
    if English == False:
        sentence = "Die Auswahl der Zelle erfolgt mittels einer Regressionsanalyse, wobei die Zelle mit der besten Korrelation für den Langzeitbezug " \
                    "verwendet wird. Die Koordinaten der gewählten Zelle sind: {}, {}m über Grund. Mittels einer Regressionsgleichung erfolgt die " \
                    "Anpassung der {}-Index-Daten an die Messzeitreihe.".format(addVar1, addVar2, addVar3)

    return sentence


def wbReader(folderIn, data, refName, weaNo):
    valList = []
    '''to read xlsm file we need openpyxl library. it has load_workbook module which helps reading the file.
        then we define the sheet name. once we are in that sheet we can use cell coordinates to access that value.'''
    wb = load_workbook("{}/{}".format(folderIn, data['files']['Positionsdaten']))
    sh = wb["Positionsdaten"]
    valList.append(str(sh["b1"].value))

    '''Production index value in Percentage'''
    wb = xlrd.open_workbook("{}/{}/wea{}.xls".format(folderIn, refName, weaNo))
    sheet = wb.sheet_by_index(0)
    a1 = "{:.1f}".format(sheet.cell_value(rowx=12, colx=8) * 100)
    valList.append(str(a1))

    '''Mean energy production in Mega-Watt'''
    with open("{}/{}/weas_{}_k2.txt".format(folderIn, refName, refName)) as fi:
        newF = fi.readlines()
    newString = newF[1].split(';')
    newInt = float(newString[2])/1000
    newInt = "{:.0f}".format(newInt)
    valList.append(str(newInt))

    return valList


def write_bookmarks(doc,bookmarks_list):
    for b in bookmarks_list:
        if b['type'] == "string":
            bookmark_name = b['bookmark']
            text = b['value']
            underline = False
            italic = False
            bold = False
            style = False
            if b.get('underline') is not None:
                underline = b['underline']
            if b.get('italic') is not None:
                italic = b['italic']
            if b.get('bold') is not None:
                bold = b['bold']
            if b.get('style') is not None:
                style = b['style']
            bookmark_text(doc, bookmark_name, text, underline, italic, bold, style)

        if b['type'] == "picture":
            bookmark_name = b['bookmark']
            path = b['value']
            bookmark_picture(doc, bookmark_name, path)
